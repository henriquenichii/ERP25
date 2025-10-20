import re
import PyPDF2
import openpyxl
from io import BytesIO
from docx import Document # Importa a biblioteca para trabalhar com .docx
from docx.shared import Inches # Para unidades de medida em .docx
import datetime # Para manipulação de datas

# ==============================================================================
# FUNÇÃO 1: Extrai o texto bruto do PDF
# ==============================================================================
def extrair_texto_de_pdf(caminho_do_pdf):
    """
    Extrai o texto de todas as páginas de um arquivo PDF.
    """
    texto_completo = ""
    try:
        with open(caminho_do_pdf, 'rb') as arquivo:
            leitor_pdf = PyPDF2.PdfReader(arquivo)
            for pagina in leitor_pdf.pages:
                texto_completo += pagina.extract_text() + "\n"
    except FileNotFoundError:
        print(f"ERRO: Arquivo não encontrado em '{caminho_do_pdf}'. Verifique o caminho.")
        return None
    except Exception as e:
        print(f"Ocorreu um erro inesperado ao ler o PDF: {e}")
        return None
    return texto_completo

# ==============================================================================
# FUNÇÃO 2: Extrai os dados específicos do texto
# ==============================================================================
def extrair_dados_do_contrato(texto_do_contrato):
    """
    Extrai dados específicos do conteúdo textual de um contrato.
    """
    dados_extraidos = {}
    flags = re.DOTALL | re.IGNORECASE

    # 1. Extrair dados do CONTRATANTE
    contratante_match = re.search(
        r"CONTRATANTE\s*:\s*Sr\(a\)\s*(.*?),\s*brasileiro\(a\).*?RG:\s*([\d.\s-]+?)\s*e\s*CPF:\s*([\d.\s-]+?),",
        texto_do_contrato, flags)
    if contratante_match:
        dados_extraidos['Contratante'] = {'Nome': contratante_match.group(1).strip(), 'RG': contratante_match.group(2).strip(), 'CPF': contratante_match.group(3).strip(),}
    else:
        dados_extraidos['Contratante'] = "Não encontrado"

    # 2. Extrair dados do CONTRATADO
    contratado_match = re.search(
        r"CONTRATADO\s*(.*?),\s*inscrito\s*sob\s*o\s*CNPJ:\s*([\d./\s-]+?),",
        texto_do_contrato, flags)
    if contratado_match:
        dados_extraidos['Contratado'] = {'Nome Empresa': contratado_match.group(1).strip(), 'CNPJ': contratado_match.group(2).strip(),}
    else:
        dados_extraidos['Contratado'] = "Não encontrado"

    # 3. Extrair Produtos Contratados
    dados_extraidos['Produtos Contratados'] = []
    secao_produtos_match = re.search(r'PRODUTOS CONTRATADOS\s*(.*?)\s*CLÁUSULA 2', texto_do_contrato, flags)
    if secao_produtos_match:
        texto_tabela = secao_produtos_match.group(1)
        linhas = texto_tabela.strip().split('\n')
        for linha in linhas:
            linha_limpa = linha.strip()
            if not linha_limpa or not linha_limpa[0].isdigit():
                continue
            # Ajuste a regex se o formato da tabela for diferente
            # Ex: (Quantidade) (Produto) (Valor Unitário) (Valor Total Item)
            produto_match = re.match(r'^\s*(\d+)\s+(.*?)\s+([\d,.]+)\s+([\d,.]+)\s*$', linha_limpa)
            if produto_match:
                dados_extraidos['Produtos Contratados'].append({
                    'Quantidade': produto_match.group(1).strip(),
                    'Produto': produto_match.group(2).strip().strip(),
                    'Valor Unitário': produto_match.group(3).strip(),
                    'Valor Total Item': produto_match.group(4).strip()
                })

    # 4, 5, 6. Extrair outros dados
    valor_total_match = re.search(r"O valor total de R\$\s*([\d,.]+)", texto_do_contrato, flags)
    dados_extraidos['Valor Total do Pedido'] = valor_total_match.group(1).strip() if valor_total_match else "Não encontrado"
    
    data_pagamento_match = re.search(r"pagos no dia\s*(\d{2}/\d{2}/\d{4})", texto_do_contrato, flags)
    dados_extraidos['Data de Pagamento'] = data_pagamento_match.group(1).strip() if data_pagamento_match else "Não encontrada"
    
    data_evento_match = re.search(r"O evento acontecerá no dia:\s*([\d/]+)", texto_do_contrato, flags)
    dados_extraidos['Data do Evento'] = data_evento_match.group(1).strip() if data_evento_match else "Não encontrada"
    
    local_evento_match = re.search(r"Local do\s*evento\s*:\s*(.*?)\n", texto_do_contrato, flags)
    dados_extraidos['Local do Evento'] = local_evento_match.group(1).strip() if local_evento_match else "Não encontrado"
            
    return dados_extraidos

# ==============================================================================
# FUNÇÃO 3: Exporta os dados para um arquivo Excel
# ==============================================================================
def exportar_para_excel(dados, nome_do_arquivo="dados_contrato.xlsx"):
    """
    Exporta o dicionário de dados extraídos para uma planilha Excel.
    Esta função é focada em dados de contrato (singular).
    Para exportar uma LISTA de pedidos, a lógica precisará ser adaptada.
    """
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Dados do Contrato"

    sheet['A1'] = "Campo"
    sheet['B1'] = "Informação Extraída"
    
    linha_atual = 2

    for chave, valor in dados.items():
        if chave == 'Produtos Contratados':
            continue
        
        if isinstance(valor, dict):
            for sub_chave, sub_valor in valor.items():
                sheet[f'A{linha_atual}'] = f"{chave} - {sub_chave}"
                sheet[f'B{linha_atual}'] = sub_valor
                linha_atual += 1
        else:
            sheet[f'A{linha_atual}'] = chave
            sheet[f'B{linha_atual}'] = valor
            linha_atual += 1

    # Adiciona a tabela de produtos contratados
    linha_atual += 2 
    
    headers_produtos = list(dados['Produtos Contratados'][0].keys()) if dados.get('Produtos Contratados') and dados['Produtos Contratados'] else []
    if headers_produtos:
        for col_idx, header in enumerate(headers_produtos, 1):
            sheet.cell(row=linha_atual, column=col_idx, value=header)
        
        linha_atual += 1

        for produto in dados['Produtos Contratados']:
            for col_idx, header in enumerate(headers_produtos, 1):
                sheet.cell(row=linha_atual, column=col_idx, value=produto.get(header, 'N/A'))
            linha_atual += 1
    
    try:
        workbook.save(filename=nome_do_arquivo)
        print(f"\n[SUCESSO] Dados exportados para o arquivo '{nome_do_arquivo}'")
    except Exception as e:
        print(f"\n[ERRO] Não foi possível salvar a planilha: {e}")

# ==============================================================================
# FUNÇÃO 4: Gera o Relatório de Entrega em DOCX (MOVIMENTO DO relatorio_entrega.py)
# ==============================================================================
def gerar_relatorio_entrega(dados, nome_arquivo='relatorio_entrega.docx'):
    """
    Gera um relatório de entrega em formato .docx com base nos dados do pedido.
    """
    document = Document()
    
    # Título
    document.add_heading('RELATÓRIO DE ENTREGA', 0)
    
    # Informações principais
    contratante = dados.get('Contratante', {})
    data_evento = dados.get('Data do Evento', 'Não informada')
    local_evento = dados.get('Local do Evento', 'Não informado')
    
    document.add_paragraph(f"Nome do Cliente: {contratante.get('Nome', 'Não encontrado')}")
    document.add_paragraph(f"Data do Evento: {data_evento}")
    document.add_paragraph(f"Local do Evento: {local_evento}")
    document.add_paragraph(f"Data de Emissão: {datetime.datetime.now().strftime('%d/%m/%Y')}")

    document.add_paragraph("\nProdutos Contratados:")
    
    # Tabela de produtos
    produtos = dados.get('Produtos Contratados', [])
    if produtos:
        tabela = document.add_table(rows=1, cols=4)
        tabela.style = 'Table Grid'
        hdr_cells = tabela.rows[0].cells
        hdr_cells[0].text = 'Quantidade'
        hdr_cells[1].text = 'Produto'
        hdr_cells[2].text = 'Valor Unitário'
        hdr_cells[3].text = 'Valor Total'

        for item in produtos:
            row_cells = tabela.add_row().cells
            row_cells[0].text = str(item.get('Quantidade', 'N/A')) # Garante string
            row_cells[1].text = str(item.get('Produto', 'N/A'))
            row_cells[2].text = str(item.get('Valor Unitário', 'N/A'))
            row_cells[3].text = str(item.get('Valor Total Item', 'N/A'))
    else:
        document.add_paragraph("Nenhum produto encontrado.")

    document.add_paragraph(f"\nValor Total do Pedido: R$ {dados.get('Valor Total do Pedido', 'Não encontrado')}")

    # Espaço para assinaturas
    document.add_paragraph("\n\n\nAssinaturas:\n")
    document.add_paragraph("______________________________\nResponsável pela Entrega")
    document.add_paragraph("\n\n")
    document.add_paragraph("______________________________\nResponsável pela Retirada")

    # Salvar o arquivo
    try:
        document.save(nome_arquivo)
        print(f"[OK] Relatório de entrega salvo em: {nome_arquivo}")
    except Exception as e:
        print(f"[ERRO] Falha ao salvar relatório de entrega: {e}")


# ==============================================================================
# FUNÇÃO 5: Gera o Contrato completo em DOCX
# ==============================================================================
def gerar_contrato_docx(dados, nome_arquivo='contrato.docx'):
    """
    Gera um contrato completo em formato .docx com base nos dados do formulário.
    A função espera um dicionário 'dados' com as seguintes chaves:
    'Contratante', 'Contratado', 'Data do Evento', 'Local do Evento',
    'Produtos Contratados', 'Valor Total do Pedido', 'Data de Pagamento', 'Forma de Pagamento',
    'Como nos conheceu', 'Responsavel'
    """
    document = Document()
    
    # Adiciona o título e informações fixas do contrato
    document.add_heading('Divinos Doces Finos', 0)
    document.add_paragraph('CONTRATO', style='Normal')
    document.add_paragraph()

    # Informações do Contratante
    contratante = dados.get('Contratante', {})
    contratante_texto = document.add_paragraph()
    contratante_texto.add_run('CONTRATANTE: ').bold = True
    contratante_texto.add_run(f"Sr(a) {contratante.get('Nome', 'N/A')}, brasileiro(a), portador(a) da cédula de RG: {contratante.get('RG', 'N/A')} e CPF: {contratante.get('CPF', 'N/A')}, residente e domiciliado(a) na {contratante.get('Endereco', 'N/A')} - Tel. {contratante.get('Telefone', 'N/A')}.")

    # Informações do Contratado (Padronizadas)
    contratado_texto = document.add_paragraph()
    contratado_texto.add_run('CONTRATADO: ').bold = True
    contratado_texto.add_run(f"Divinos Doces Finos, inscrito sob o CNPJ: 18.826.801/0001-76, com sede na Rua Curupacê, 392 Mooca, São Paulo SP representado pela sócia proprietária Damaris Talita Macedo, portador do RG: 30.315.655-7.")
    document.add_paragraph()
    
    # CLÁUSULA 1 - PRODUTOS CONTRATADOS
    document.add_heading('CLÁUSULA 1 - PRODUTOS CONTRATADOS', level=1)
    
    produtos = dados.get('Produtos Contratados', [])
    if produtos:
        tabela = document.add_table(rows=1, cols=4)
        tabela.style = 'Table Grid'
        hdr_cells = tabela.rows[0].cells
        hdr_cells[0].text = 'Quantidade'
        hdr_cells[1].text = 'Produto'
        hdr_cells[2].text = 'Valor Unitário'
        hdr_cells[3].text = 'Valor Total'
        
        for item in produtos:
            row_cells = tabela.add_row().cells
            row_cells[0].text = str(item.get('Quantidade', 'N/A'))
            row_cells[1].text = str(item.get('Produto', 'N/A'))
            row_cells[2].text = str(item.get('Valor Unitário', 'N/A'))
            row_cells[3].text = str(item.get('Valor Total Item', 'N/A'))
        document.add_paragraph(f"TOTAL: R$ {dados.get('Valor Total do Pedido', 'N/A')}")
    else:
        document.add_paragraph("Nenhum produto adicionado.")

    # CLÁUSULA 2 - VALOR E FORMA DE PAGAMENTO
    document.add_heading('CLÁUSULA 2 - VALOR E FORMA DE PAGAMENTO', level=1)
    document.add_paragraph(f"O valor total de R$ {dados.get('Valor Total do Pedido', 'N/A')} referente aos produtos acima citados, foram pagos no dia {dados.get('Data de Pagamento', 'N/A')} {dados.get('Forma de Pagamento', 'N/A')}.")

    # ... (Adicionar as outras cláusulas de forma fixa, como no seu modelo) ...
    document.add_heading('CLÁUSULA 3 - EMBALAGEM DOS DOCES - FORMINHAS', level=1)
    document.add_paragraph('Os doces finos são entregues em forminhas no formato caixeta, na cor branca, todos decorados e prontos para o consumo. Os brigadeiros serão entregues em forminhas na cor branca nº 5.')
    document.add_paragraph('Caso o CONTRATANTE opte por embalagens decorativas, o mesmo deverá enviar ao CONTRATADO com no máximo 15 dias de antecedência ao evento, que entregará os doces finos dentro das embalagens decoradas, prontos para o consumo. Após esse prazo não recebemos.')
    document.add_paragraph('Por haver um manejo especial nas forminhas no modelo de flor e um custo maior de compra de caixas para armazenamento dos doces, é cobrado uma taxa adicional de R$0,10 por unidade, como consta abaixo:')
    document.add_paragraph('ATÉ 100 DOCES + R$10,00 / ATÉ 200 DOCES + R$20,00 ATÉ 300 DOCES + R$30,00 / ATÉ 400 DOCES + R$40,00 ACIMA DE 500 DOCES + R$50,00 e assim sucessivamente')
    document.add_paragraph()

    # CLÁUSULA 4 - EMBALAGENS DOS BEM-CASADOS
    document.add_heading('CLÁUSULA 4 - EMBALAGENS DOS BEM-CASADOS', level=1)
    document.add_paragraph('Os bem-casados são entregues em papel crepom crepe plus, com celofane e fita de cetim de 7mm, nas cores enviadas na tabela completa. Os papéis perolados da linha especial serão cobrados R$ 0,40 a mais por unidade e os papéis dourado, prata, tiffany e marsala serão cobrados R$ 0,20 a mais por unidade, por se tratar de um papel especial e com maior custo. Tudo está discriminado na tabela de cores.')
    document.add_paragraph('Caso o CONTRATANTE opte por incluir, medalhinhas, tercinhos, renda, juta, tag ou outro item decorativo, deverá consultar antecipadamente a disponibilidade e todos os itens são colados com cola quente. A entrega dos itens deverá ocorrer com no máximo 15 dias antes do evento. Após esse prazo não recebemos. Por haver um manejo especial dos itens, será cobrado uma taxa adicional de R$0,10, como consta abaixo: ATÉ 100 BEM-CASADOS + R$10,00 / ATÉ 200 BEM-CASADOS + R$20,00 ATÉ 300 BEM-CASADOS + R$30,00 / ATÉ 400 BEM-CASADOS + R$40,00 ACIMA DE 500 BEM-CASADOS + R$50,00 e assim sucessivamente')
    document.add_paragraph('Caso opte pela aplicação de dois ou mais itens, será cobrado o valor de cada item.')
    document.add_paragraph()

    # CLÁUSULA 5 - ALTERAÇÕES
    document.add_heading('CLÁUSULA 5 - ALTERAÇÕES', level=1)
    document.add_paragraph('Não recebemos forminhas, modificações, alterações em contrato em hipótese alguma na semana do evento.')
    document.add_paragraph()
    
    # CLÁUSULA 6 - ADIÇÃO DE NOVOS ITENS
    document.add_heading('CLÁUSULA 6 - ADIÇÃO DE NOVOS ITENS', level=1)
    document.add_paragraph('Caso haja a necessidade do CONTRATANTE adicionar novos itens ao pedido fechado, o valor dos produtos será de acordo com o valor vigente no momento da adição, mesmo que o contrato tenha sido fechado com valores promocionais.')
    document.add_paragraph('A adição de produtos ocorre de acordo com a disponibilidade de agenda. Não havendo disponibilidade para novos produtos ou pedidos, não será possível a complementação.')
    document.add_paragraph()

    # CLÁUSULA 7 - RETIRADA OU SERVIÇO DE ENTREGA
    document.add_heading('CLÁUSULA 7 - RETIRADA OU SERVIÇO DE ENTREGA', level=1)
    document.add_paragraph(f"A entrega ou retirada dos itens acima, deverá ser definida pela CONTRATANTE até 15 dias antes do evento. Em caso de entrega será cobrada taxa de deslocamento de R$ 6,00 por km ou a taxa mínima de R$50,00 (sujeito a disponibilidade na data e horário desejados). Não fazemos entregas aos domingos e feriados. A retirada dos produtos ocorre de segunda-feira à sábado, das 9h às 16h30, mediante agendamento com o setor responsável, não havendo expediente aos domingos e feriados.")
    document.add_paragraph()
    
    # CLÁUSULA 8 - ARMAZENAMENTO
    document.add_heading('CLÁUSULA 8 - ARMAZENAMENTO', level=1)
    document.add_paragraph('Todos os doces e/ou bem-casados deverão, obrigatoriamente, ser armazenados em geladeira até o momento da montagem da mesa para o evento. Validade 3 a 5 dias em geladeira.')
    document.add_paragraph()

    # CLÁUSULA 9 - LOCAÇÃO (SE HOUVER)
    document.add_heading('CLÁUSULA 9 - LOCAÇÃO (SE HOUVER)', level=1)
    document.add_paragraph('Caso haja locação de bolo cenográfico, o CONTRATANTE deverá deixar uma caução no valor de R$300,00 ou o valor em dinheiro, como forma de garantia. O bolo cenográfico sendo locado e deverá retornar nas mesmas condições, em até 4 dias após a data da retirada. Na devolução do bolo cenográfico, será devolvido o valor total. Em caso de avarias será cobrado R$ 100,00 por andar (dependendo do modelo) para refazer cada andar danificado. O CONTRATANTE deverá tomar todos os cuidados necessários como: não expor ao calor excessivo, água ou qualquer outro líquido, não deverá apertar, amassar, não deixar convidados colocarem as mãos e deverá ser transportado com cuidado, pegando somente pela base de madeira.')
    document.add_paragraph()

    # CLÁUSULA 10 - REMARCAÇÃO
    document.add_heading('CLÁUSULA 10 - REMARCAÇÃO', level=1)
    document.add_paragraph('Em caso de REMARCAÇÃO de data do evento superior a 6 meses, será cobrado um reequilíbrio econômico e financeiro de 10% sobre o valor do contrato, a cada 6 meses de diferença da data marcada inicialmente.')
    document.add_paragraph()

    # CLÁUSULA 11 DATA E LOCAL DO EVENTO
    document.add_heading('CLÁUSULA 11 - DATA E LOCAL DO EVENTO', level=1)
    document.add_paragraph(f"O evento acontecerá no dia: {dados.get('Data do Evento', 'N/A')} - Local do evento: {dados.get('Local do Evento', 'N/A')}")
    document.add_paragraph(f"Como nos conheceu: {dados.get('Como nos conheceu', 'N/A')}")
    document.add_paragraph()

    # CLÁUSULA 12 - CANCELAMENTO
    document.add_heading('CLÁUSULA 12 - CANCELAMENTO', level=1)
    document.add_paragraph('A CONTRATANTE pagará multa de 30% do valor do contrato em caso de cancelamento. O CONTRATADO pagará multa de 100% do valor do contrato em caso de cancelamento.')
    document.add_paragraph()

    # Assinaturas
    document.add_paragraph(f"RESPONSÁVEL PELO CONTRATO: {dados.get('Responsavel', 'N/A')}")
    document.add_paragraph(f"São Paulo, {dados.get('Data de Pagamento', 'N/A')}")
    document.add_paragraph()
    document.add_paragraph('CONTRATANTE', style='Normal').bold = True
    document.add_paragraph('______________________________', style='Normal')
    document.add_paragraph('CONTRATADO', style='Normal').bold = True
    document.add_paragraph('______________________________', style='Normal')

    # Salvar o arquivo
    try:
        doc_stream = BytesIO()
        document.save(doc_stream)
        doc_stream.seek(0)
        print("[OK] Contrato gerado na memória.")
        return doc_stream
    except Exception as e:
        print(f"[ERRO] Falha ao gerar contrato: {e}")
        return None


# ==============================================================================
# BLOCO DE EXECUÇÃO: O "Gerente de Operações" para testes local
# ==============================================================================
if __name__ == "__main__":
    caminho_do_pdf = "modelo_contrato.pdf" # Certifique-se de ter este arquivo no mesmo diretório

    # Exemplo de uso das funções
    print(f"Tentando extrair texto de: {caminho_do_pdf}")
    texto_extraido = extrair_texto_de_pdf(caminho_do_pdf)
    
    if texto_extraido:
        print("\nTexto extraído com sucesso. Tentando extrair dados...")
        dados_do_contrato = extrair_dados_do_contrato(texto_extraido)
        
        # Imprime os resultados no terminal
        print("\n--- DADOS EXTRAÍDOS DO CONTRATO ---")
        for chave, valor in dados_do_contrato.items():
            print(f"\n>> {chave}:")
            if isinstance(valor, dict):
                for sub_chave, sub_valor in valor.items():
                    print(f"    {sub_chave}: {sub_valor}")
            elif isinstance(valor, list) and valor:
                for item in valor:
                    print(f"    - {item}")
            else:
                print(f"    {valor}")
        print("\n----------------------------------------------------")

        # Chama a função para criar a planilha Excel
        nome_excel = f"dados_contrato_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        exportar_para_excel(dados_do_contrato, nome_do_arquivo=nome_excel)

        # Chama a nova função para criar o relatório de entrega em DOCX
        nome_docx = f"relatorio_entrega_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        gerar_relatorio_entrega(dados_do_contrato, nome_arquivo=nome_docx)

    else:
        print("\nNão foi possível processar o PDF.")